"""
Document Exporter Module
Advanced document export with DOCX, PDF, HTML, Markdown support.
Uses: python-docx, reportlab, weasyprint, markdown, aiofiles
"""

import asyncio
import logging
from pathlib import Path
from typing import Any, Optional, Dict, List
from datetime import datetime
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)

# Optional imports - graceful degradation
try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    AIOFILES_AVAILABLE = False
    logger.warning("aiofiles not available. Using sync file operations.")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX export disabled.")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab not available. PDF export disabled.")

try:
    import markdown as md
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logger.warning("markdown not available. Markdown conversion disabled.")

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logger.warning("weasyprint not available. HTML to PDF disabled.")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl not available. Advanced Excel disabled.")


class ExportFormat(Enum):
    """Supported export formats."""
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    MARKDOWN = "md"
    EXCEL = "xlsx"
    JSON = "json"
    TEXT = "txt"


@dataclass
class ExportResult:
    """Result of document export."""
    success: bool
    file_path: str
    format: str
    size_bytes: int
    message: str


class DocumentExporter:
    """
    Advanced document exporter supporting multiple formats.
    """
    
    def __init__(self, output_dir: str = "./exports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"DocumentExporter initialized. Output: {self.output_dir}")
        self._log_available_formats()
    
    def _log_available_formats(self):
        """Log available export formats."""
        formats = []
        if DOCX_AVAILABLE:
            formats.append("DOCX")
        if REPORTLAB_AVAILABLE:
            formats.append("PDF")
        if MARKDOWN_AVAILABLE:
            formats.append("Markdown")
        if WEASYPRINT_AVAILABLE:
            formats.append("HTML→PDF")
        if OPENPYXL_AVAILABLE:
            formats.append("Excel")
        logger.info(f"Available export formats: {', '.join(formats) or 'Text only'}")
    
    async def export(
        self,
        content: Any,
        filename: str,
        format: ExportFormat,
        title: str = "Document",
        metadata: Dict[str, Any] = None
    ) -> ExportResult:
        """Export content to specified format."""
        metadata = metadata or {}
        
        try:
            if format == ExportFormat.DOCX:
                return await self._export_docx(content, filename, title, metadata)
            elif format == ExportFormat.PDF:
                return await self._export_pdf(content, filename, title, metadata)
            elif format == ExportFormat.HTML:
                return await self._export_html(content, filename, title, metadata)
            elif format == ExportFormat.MARKDOWN:
                return await self._export_markdown(content, filename, title, metadata)
            elif format == ExportFormat.EXCEL:
                return await self._export_excel(content, filename, title, metadata)
            elif format == ExportFormat.JSON:
                return await self._export_json(content, filename, metadata)
            else:
                return await self._export_text(content, filename, metadata)
        
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return ExportResult(
                success=False,
                file_path="",
                format=format.value,
                size_bytes=0,
                message=str(e)
            )
    
    async def _export_docx(
        self,
        content: Any,
        filename: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for DOCX export")
        
        file_path = self.output_dir / f"{filename}.docx"
        
        doc = Document()
        
        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        if metadata.get("author"):
            doc.add_paragraph(f"Author: {metadata['author']}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph()
        
        # Content
        if isinstance(content, str):
            for line in content.split('\n'):
                if line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                elif line.strip():
                    doc.add_paragraph(line)
        elif isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item), style='List Bullet')
        elif isinstance(content, dict):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Key'
            hdr_cells[1].text = 'Value'
            for key, value in content.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value)
        
        doc.save(str(file_path))
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="docx",
            size_bytes=size,
            message="DOCX exported successfully"
        )
    
    async def _export_pdf(
        self,
        content: Any,
        filename: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to PDF using reportlab."""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab required for PDF export")
        
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.platypus import PageBreak
        
        file_path = self.output_dir / f"{filename}.pdf"
        
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        # Helper to escape XML special chars
        def escape_xml(text: str) -> str:
            if not text:
                return ""
            return (str(text)
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;'))
        
        # Title
        story.append(Paragraph(escape_xml(title), title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        story.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            styles['Normal']
        ))
        story.append(Spacer(1, 24))
        
        # Content
        if isinstance(content, str) and content.strip():
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue
                
                escaped_line = escape_xml(line)
                
                if line.startswith('# '):
                    story.append(Paragraph(escape_xml(line[2:]), styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(escape_xml(line[3:]), styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(escape_xml(line[4:]), styles['Heading3']))
                elif line.startswith('- ') or line.startswith('* '):
                    # Bullet points
                    bullet_text = f"• {escape_xml(line[2:])}"
                    story.append(Paragraph(bullet_text, body_style))
                else:
                    story.append(Paragraph(escaped_line, body_style))
                story.append(Spacer(1, 4))
                
        elif isinstance(content, list) and content:
            for item in content:
                if isinstance(item, dict):
                    # Each dict item as a mini-section
                    for k, v in item.items():
                        story.append(Paragraph(f"<b>{escape_xml(str(k))}:</b> {escape_xml(str(v))}", body_style))
                    story.append(Spacer(1, 10))
                else:
                    story.append(Paragraph(f"• {escape_xml(str(item))}", body_style))
                    story.append(Spacer(1, 4))
                    
        elif isinstance(content, dict) and content:
            data = [['Key', 'Value']]
            for k, v in content.items():
                data.append([escape_xml(str(k)), escape_xml(str(v))])
            
            table = Table(data, colWidths=[150, 300])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f5f5f5')),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(table)
        else:
            # Empty or unsupported content - add placeholder
            story.append(Paragraph("No content to display.", body_style))
        
        # Ensure there's at least some content
        if len(story) < 3:
            story.append(Spacer(1, 50))
            story.append(Paragraph("Document generated successfully.", body_style))
        
        doc.build(story)
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="pdf",
            size_bytes=size,
            message="PDF exported successfully"
        )
    
    async def _export_html(
        self,
        content: Any,
        filename: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to HTML."""
        file_path = self.output_dir / f"{filename}.html"
        
        # Convert markdown to HTML if available
        if isinstance(content, str) and MARKDOWN_AVAILABLE:
            html_content = md.markdown(content, extensions=['tables', 'fenced_code'])
        else:
            html_content = f"<pre>{content}</pre>"
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 2rem; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
        pre {{ background: #f4f4f4; padding: 1rem; border-radius: 4px; overflow-x: auto; }}
        code {{ background: #f4f4f4; padding: 0.2rem 0.4rem; border-radius: 3px; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background: #f2f2f2; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p><small>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</small></p>
    <hr>
    {html_content}
</body>
</html>"""
        
        await self._write_file(file_path, html)
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="html",
            size_bytes=size,
            message="HTML exported successfully"
        )
    
    async def _export_markdown(
        self,
        content: Any,
        filename: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to Markdown."""
        file_path = self.output_dir / f"{filename}.md"
        
        lines = [
            f"# {title}",
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
            ""
        ]
        
        if isinstance(content, str):
            lines.append(content)
        elif isinstance(content, dict):
            lines.append("| Key | Value |")
            lines.append("|-----|-------|")
            for k, v in content.items():
                lines.append(f"| {k} | {v} |")
        elif isinstance(content, list):
            for item in content:
                lines.append(f"- {item}")
        
        md_content = '\n'.join(lines)
        await self._write_file(file_path, md_content)
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="md",
            size_bytes=size,
            message="Markdown exported successfully"
        )
    
    async def _export_excel(
        self,
        content: Any,
        filename: str,
        title: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to Excel using openpyxl."""
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl required for Excel export")
        
        file_path = self.output_dir / f"{filename}.xlsx"
        
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]  # Excel sheet name limit
        
        # Header styling
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        if isinstance(content, dict):
            ws.append(["Key", "Value"])
            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
            for k, v in content.items():
                ws.append([str(k), str(v)])
        elif isinstance(content, list):
            if content and isinstance(content[0], dict):
                headers = list(content[0].keys())
                ws.append(headers)
                for cell in ws[1]:
                    cell.font = header_font
                    cell.fill = header_fill
                for item in content:
                    ws.append([item.get(h, "") for h in headers])
            else:
                ws.append(["Value"])
                ws[1][0].font = header_font
                ws[1][0].fill = header_fill
                for item in content:
                    ws.append([str(item)])
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[column_letter].width = min(max_length + 2, 50)
        
        wb.save(str(file_path))
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="xlsx",
            size_bytes=size,
            message="Excel exported successfully"
        )
    
    async def _export_json(
        self,
        content: Any,
        filename: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to JSON."""
        import json
        file_path = self.output_dir / f"{filename}.json"
        
        data = {
            "generated": datetime.now().isoformat(),
            "metadata": metadata,
            "content": content
        }
        
        await self._write_file(file_path, json.dumps(data, indent=2, default=str))
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="json",
            size_bytes=size,
            message="JSON exported successfully"
        )
    
    async def _export_text(
        self,
        content: Any,
        filename: str,
        metadata: Dict[str, Any]
    ) -> ExportResult:
        """Export to plain text."""
        file_path = self.output_dir / f"{filename}.txt"
        
        await self._write_file(file_path, str(content))
        size = file_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(file_path),
            format="txt",
            size_bytes=size,
            message="Text exported successfully"
        )
    
    async def _write_file(self, path: Path, content: str) -> None:
        """Write file using aiofiles if available."""
        if AIOFILES_AVAILABLE:
            async with aiofiles.open(path, 'w', encoding='utf-8') as f:
                await f.write(content)
        else:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
    
    async def html_to_pdf(self, html_path: str, pdf_path: str = None) -> ExportResult:
        """Convert HTML to PDF using weasyprint."""
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("weasyprint required for HTML to PDF conversion")
        
        html_file = Path(html_path)
        if not pdf_path:
            pdf_path = html_file.with_suffix('.pdf')
        else:
            pdf_path = Path(pdf_path)
        
        HTML(filename=str(html_file)).write_pdf(str(pdf_path))
        size = pdf_path.stat().st_size
        
        return ExportResult(
            success=True,
            file_path=str(pdf_path),
            format="pdf",
            size_bytes=size,
            message="HTML converted to PDF successfully"
        )
    
    def get_available_formats(self) -> List[str]:
        """Get list of available export formats."""
        formats = ["json", "txt", "html", "md"]
        if DOCX_AVAILABLE:
            formats.append("docx")
        if REPORTLAB_AVAILABLE:
            formats.append("pdf")
        if OPENPYXL_AVAILABLE:
            formats.append("xlsx")
        return formats
