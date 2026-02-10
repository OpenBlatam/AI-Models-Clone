"""Word Converter - Convert Markdown to Word format"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, Any, Optional

from .base_converter import BaseConverter


class WordConverter(BaseConverter):
    """Convert Markdown to Word (.docx) format"""
    
    async def convert(
        self,
        parsed_content: Dict[str, Any],
        output_path: str,
        include_charts: bool = True,
        include_tables: bool = True,
        custom_styling: Optional[Dict[str, Any]] = None
    ) -> None:
        """Convert to Word format"""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add title
        if parsed_content.get("title"):
            title = doc.add_heading(parsed_content["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.size = Pt(24)
            title.runs[0].font.color.rgb = RGBColor(26, 26, 26)
        
        # Add headings
        for heading in parsed_content.get("headings", []):
            level = min(heading["level"], 9)  # Word supports up to level 9
            doc.add_heading(heading["text"], level=level)
        
        # Add tables
        if include_tables:
            for table_data in parsed_content.get("tables", []):
                # Create table
                table = doc.add_table(rows=1, cols=len(table_data["headers"]))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                header_cells = table.rows[0].cells
                for idx, header in enumerate(table_data["headers"]):
                    header_cells[idx].text = header
                    header_cells[idx].paragraphs[0].runs[0].font.bold = True
                
                # Add data rows
                for row_data in table_data["rows"]:
                    row_cells = table.add_row().cells
                    for idx, value in enumerate(row_data):
                        if idx < len(row_cells):
                            row_cells[idx].text = str(value)
                
                doc.add_paragraph()  # Add spacing
        
        # Add paragraphs
        for para in parsed_content.get("paragraphs", []):
            doc.add_paragraph(para)
        
        # Add lists
        for list_data in parsed_content.get("lists", []):
            if list_data["type"] == "unordered":
                for item in list_data["items"]:
                    doc.add_paragraph(item, style='List Bullet')
            else:
                for item in list_data["items"]:
                    doc.add_paragraph(item, style='List Number')
        
        # Save document
        doc.save(output_path)

