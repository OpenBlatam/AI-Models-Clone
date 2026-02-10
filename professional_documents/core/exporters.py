"""
Document export formatters.

Abstract base class and implementations for different export formats.
"""

import aiofiles
from abc import ABC, abstractmethod
from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import ProfessionalDocument, DocumentExportRequest
from .export_helpers import (
    format_date,
    get_document_metadata,
    format_section_heading,
    has_metadata,
    build_markdown_metadata_lines,
    build_text_metadata_lines
)


class DocumentExporter(ABC):
    """Abstract base class for document exporters."""
    
    @abstractmethod
    async def export(self, document: ProfessionalDocument, request: DocumentExportRequest, output_path: Path) -> Path:
        """Export document to the specified format."""
        pass
    
    @abstractmethod
    def get_file_extension(self) -> str:
        """Get the file extension for this export format."""
        pass


class PDFExporter(DocumentExporter):
    """PDF document exporter."""
    
    def get_file_extension(self) -> str:
        return "pdf"
    
    async def export(self, document: ProfessionalDocument, request: DocumentExportRequest, output_path: Path) -> Path:
        """Export document to PDF format."""
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=document.style.margin_right * inch,
            leftMargin=document.style.margin_left * inch,
            topMargin=document.style.margin_top * inch,
            bottomMargin=document.style.margin_bottom * inch
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=colors.HexColor(document.style.header_color),
            fontName=document.style.font_family
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor(document.style.header_color),
            fontName=document.style.font_family
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=document.style.font_size,
            spaceAfter=6,
            textColor=colors.HexColor(document.style.body_color),
            fontName=document.style.font_family,
            leading=document.style.font_size * document.style.line_spacing
        )
        
        story = []
        story.append(Paragraph(document.title, title_style))
        story.append(Spacer(1, 12))
        
        if document.subtitle:
            story.append(Paragraph(document.subtitle, heading_style))
            story.append(Spacer(1, 12))
        
        if has_metadata(document):
            author_info = build_text_metadata_lines(document)
            story.append(Paragraph("<br/>".join(author_info), body_style))
            story.append(Spacer(1, 12))
        story.append(Spacer(1, 20))
        
        for section in document.sections:
            if section.level == 1:
                story.append(Paragraph(section.title, title_style))
            elif section.level == 2:
                story.append(Paragraph(section.title, heading_style))
            else:
                story.append(Paragraph(section.title, body_style))
            
            story.append(Spacer(1, 6))
            story.append(Paragraph(section.content, body_style))
            story.append(Spacer(1, 12))
        
        doc.build(story)
        return output_path


class MarkdownExporter(DocumentExporter):
    """Markdown document exporter."""
    
    def get_file_extension(self) -> str:
        return "md"
    
    async def export(self, document: ProfessionalDocument, request: DocumentExportRequest, output_path: Path) -> Path:
        """Export document to Markdown format."""
        markdown_content = []
        
        markdown_content.append(f"# {document.title}")
        markdown_content.append("")
        
        if document.subtitle:
            markdown_content.append(f"## {document.subtitle}")
            markdown_content.append("")
        
        if has_metadata(document):
            markdown_content.append("---")
            markdown_content.extend(build_markdown_metadata_lines(document))
            markdown_content.append("---")
            markdown_content.append("")
        
        for section in document.sections:
            heading_level = "#" * format_section_heading(section)
            markdown_content.append(f"{heading_level} {section.title}")
            markdown_content.append("")
            markdown_content.append(section.content)
            markdown_content.append("")
        
        async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
            await f.write('\n'.join(markdown_content))
        
        return output_path


class WordExporter(DocumentExporter):
    """Word document exporter."""
    
    def get_file_extension(self) -> str:
        return "docx"
    
    async def export(self, document: ProfessionalDocument, request: DocumentExportRequest, output_path: Path) -> Path:
        """Export document to Word format."""
        doc = DocxDocument()
        
        doc.core_properties.title = document.title
        if document.author:
            doc.core_properties.author = document.author
        doc.core_properties.created = document.date_created
        
        title = doc.add_heading(document.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if document.subtitle:
            subtitle = doc.add_heading(document.subtitle, 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if has_metadata(document):
            info_para = doc.add_paragraph()
            metadata_lines = build_text_metadata_lines(document)
            info_para.add_run("\n".join(metadata_lines))
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        for section in document.sections:
            heading = doc.add_heading(section.title, section.level)
            content_para = doc.add_paragraph(section.content)
            
            for run in content_para.runs:
                run.font.name = document.style.font_family
                run.font.size = Pt(document.style.font_size)
        
        doc.save(str(output_path))
        return output_path


class HTMLExporter(DocumentExporter):
    """HTML document exporter."""
    
    def get_file_extension(self) -> str:
        return "html"
    
    async def export(self, document: ProfessionalDocument, request: DocumentExportRequest, output_path: Path) -> Path:
        """Export document to HTML format."""
        html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{document.title}</title>
    <style>
        body {{
            font-family: {document.style.font_family};
            font-size: {document.style.font_size}px;
            line-height: {document.style.line_spacing};
            color: {document.style.body_color};
            background-color: {document.style.background_color};
            margin: {document.style.margin_top}in {document.style.margin_right}in {document.style.margin_bottom}in {document.style.margin_left}in;
        }}
        h1, h2, h3, h4, h5, h6 {{
            color: {document.style.header_color};
        }}
        .title {{
            text-align: center;
            margin-bottom: 2em;
        }}
        .metadata {{
            text-align: center;
            margin-bottom: 2em;
            font-style: italic;
        }}
        .section {{
            margin-bottom: 1.5em;
        }}
    </style>
</head>
<body>
    <div class="title">
        <h1>{document.title}</h1>
        {f'<h2>{document.subtitle}</h2>' if document.subtitle else ''}
    </div>
    
    <div class="metadata">
        {f'<p><strong>Author:</strong> {document.author}</p>' if document.author else ''}
        {f'<p><strong>Company:</strong> {document.company}</p>' if document.company else ''}
        <p><strong>Date:</strong> {format_date(document.date_created)}</p>
    </div>
    
    <div class="content">
"""
        
        for section in document.sections:
            heading_tag = f"h{format_section_heading(section)}"
            html_content += f"""
        <div class="section">
            <{heading_tag}>{section.title}</{heading_tag}>
            <p>{section.content.replace(chr(10), '<br>')}</p>
        </div>
"""
        
        html_content += """
    </div>
</body>
</html>
"""
        
        async with aiofiles.open(output_path, 'w', encoding='utf-8') as f:
            await f.write(html_content)
        
        return output_path
