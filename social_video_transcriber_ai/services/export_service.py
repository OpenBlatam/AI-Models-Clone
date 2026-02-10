"""
Export Service for Social Video Transcriber AI
Exports transcriptions to multiple formats: PDF, DOCX, JSON, TXT, SRT, VTT
"""

import json
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime
from dataclasses import dataclass
from enum import Enum
import io

from ..config.settings import get_settings
from ..core.models import TranscriptionResponse, TranscriptionSegment, ContentAnalysis

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    """Supported export formats"""
    JSON = "json"
    TXT = "txt"
    SRT = "srt"
    VTT = "vtt"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "md"


@dataclass
class ExportOptions:
    """Export configuration options"""
    include_timestamps: bool = True
    include_analysis: bool = True
    include_metadata: bool = True
    include_segments: bool = True
    format_style: str = "professional"  # professional, casual, minimal
    language: str = "es"


class ExportService:
    """Service for exporting transcriptions to various formats"""
    
    def __init__(self):
        self.settings = get_settings()
        self.export_dir = Path(self.settings.temp_dir) / "exports"
        self.export_dir.mkdir(parents=True, exist_ok=True)
    
    async def export(
        self,
        transcription: TranscriptionResponse,
        format: ExportFormat,
        options: Optional[ExportOptions] = None,
    ) -> bytes:
        """
        Export transcription to specified format
        
        Args:
            transcription: Transcription response to export
            format: Export format
            options: Export options
            
        Returns:
            Exported content as bytes
        """
        options = options or ExportOptions()
        
        exporters = {
            ExportFormat.JSON: self._export_json,
            ExportFormat.TXT: self._export_txt,
            ExportFormat.SRT: self._export_srt,
            ExportFormat.VTT: self._export_vtt,
            ExportFormat.HTML: self._export_html,
            ExportFormat.MARKDOWN: self._export_markdown,
            ExportFormat.PDF: self._export_pdf,
            ExportFormat.DOCX: self._export_docx,
        }
        
        exporter = exporters.get(format)
        if not exporter:
            raise ValueError(f"Unsupported format: {format}")
        
        return await exporter(transcription, options)
    
    async def _export_json(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to JSON format"""
        data = {
            "job_id": str(transcription.job_id),
            "video_title": transcription.video_title,
            "video_author": transcription.video_author,
            "video_duration": transcription.video_duration,
            "platform": transcription.platform_detected.value if transcription.platform_detected else None,
            "created_at": transcription.created_at.isoformat(),
            "full_text": transcription.full_text,
        }
        
        if options.include_timestamps and transcription.full_text_with_timestamps:
            data["full_text_with_timestamps"] = transcription.full_text_with_timestamps
        
        if options.include_segments and transcription.segments:
            data["segments"] = [
                {
                    "id": s.id,
                    "start_time": s.start_time,
                    "end_time": s.end_time,
                    "text": s.text,
                    "formatted_timestamp": s.formatted_timestamp,
                }
                for s in transcription.segments
            ]
        
        if options.include_analysis and transcription.analysis:
            data["analysis"] = transcription.analysis.to_dict()
        
        return json.dumps(data, ensure_ascii=False, indent=2).encode('utf-8')
    
    async def _export_txt(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to plain text format"""
        lines = []
        
        if options.include_metadata:
            lines.append(f"# {transcription.video_title or 'Transcripción'}")
            lines.append(f"Autor: {transcription.video_author or 'Desconocido'}")
            if transcription.video_duration:
                minutes = int(transcription.video_duration // 60)
                seconds = int(transcription.video_duration % 60)
                lines.append(f"Duración: {minutes}:{seconds:02d}")
            lines.append(f"Fecha: {transcription.created_at.strftime('%Y-%m-%d %H:%M')}")
            lines.append("")
            lines.append("=" * 50)
            lines.append("")
        
        if options.include_timestamps:
            lines.append(transcription.full_text_with_timestamps or transcription.full_text or "")
        else:
            lines.append(transcription.full_text or "")
        
        if options.include_analysis and transcription.analysis:
            lines.append("")
            lines.append("=" * 50)
            lines.append("ANÁLISIS")
            lines.append("=" * 50)
            lines.append(f"Framework: {transcription.analysis.framework.value}")
            lines.append(f"Tono: {transcription.analysis.tone}")
            if transcription.analysis.key_points:
                lines.append("\nPuntos clave:")
                for point in transcription.analysis.key_points:
                    lines.append(f"  • {point}")
        
        return '\n'.join(lines).encode('utf-8')
    
    async def _export_srt(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to SRT subtitle format"""
        if not transcription.segments:
            return b""
        
        lines = []
        for i, segment in enumerate(transcription.segments, 1):
            start = self._format_srt_time(segment.start_time)
            end = self._format_srt_time(segment.end_time)
            lines.append(str(i))
            lines.append(f"{start} --> {end}")
            lines.append(segment.text)
            lines.append("")
        
        return '\n'.join(lines).encode('utf-8')
    
    async def _export_vtt(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to WebVTT subtitle format"""
        if not transcription.segments:
            return b"WEBVTT\n\n"
        
        lines = ["WEBVTT", ""]
        for segment in transcription.segments:
            start = self._format_vtt_time(segment.start_time)
            end = self._format_vtt_time(segment.end_time)
            lines.append(f"{start} --> {end}")
            lines.append(segment.text)
            lines.append("")
        
        return '\n'.join(lines).encode('utf-8')
    
    async def _export_html(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to HTML format"""
        html = f"""<!DOCTYPE html>
<html lang="{options.language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{transcription.video_title or 'Transcripción'}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{ color: #1a1a2e; border-bottom: 2px solid #4a4e69; padding-bottom: 10px; }}
        .metadata {{ color: #666; font-size: 0.9em; margin-bottom: 20px; }}
        .segment {{ margin: 10px 0; padding: 10px; background: #f8f9fa; border-radius: 4px; }}
        .timestamp {{ color: #4a4e69; font-weight: bold; font-size: 0.85em; }}
        .analysis {{ background: #e8f4f8; padding: 20px; border-radius: 8px; margin-top: 30px; }}
        .analysis h2 {{ color: #2c3e50; margin-top: 0; }}
        .key-points {{ list-style: none; padding: 0; }}
        .key-points li {{ padding: 5px 0; padding-left: 20px; position: relative; }}
        .key-points li::before {{ content: "→"; position: absolute; left: 0; color: #4a4e69; }}
    </style>
</head>
<body>
    <h1>{transcription.video_title or 'Transcripción de Video'}</h1>
    <div class="metadata">
        <p><strong>Autor:</strong> {transcription.video_author or 'Desconocido'}</p>
        <p><strong>Plataforma:</strong> {transcription.platform_detected.value if transcription.platform_detected else 'N/A'}</p>
        <p><strong>Duración:</strong> {self._format_duration(transcription.video_duration)}</p>
        <p><strong>Fecha:</strong> {transcription.created_at.strftime('%d/%m/%Y %H:%M')}</p>
    </div>
    
    <h2>Transcripción</h2>
"""
        
        if options.include_segments and transcription.segments:
            for segment in transcription.segments:
                html += f"""    <div class="segment">
        <span class="timestamp">{segment.formatted_timestamp}</span>
        <p>{segment.text}</p>
    </div>
"""
        else:
            html += f"    <p>{transcription.full_text}</p>\n"
        
        if options.include_analysis and transcription.analysis:
            html += f"""
    <div class="analysis">
        <h2>Análisis de Contenido</h2>
        <p><strong>Framework:</strong> {transcription.analysis.framework.value}</p>
        <p><strong>Tono:</strong> {transcription.analysis.tone}</p>
        <p><strong>Audiencia:</strong> {transcription.analysis.target_audience or 'General'}</p>
"""
            if transcription.analysis.key_points:
                html += "        <h3>Puntos Clave</h3>\n        <ul class=\"key-points\">\n"
                for point in transcription.analysis.key_points:
                    html += f"            <li>{point}</li>\n"
                html += "        </ul>\n"
            
            if transcription.analysis.hashtags_suggested:
                html += f"        <p><strong>Hashtags:</strong> {' '.join(transcription.analysis.hashtags_suggested)}</p>\n"
            
            html += "    </div>\n"
        
        html += """</body>
</html>"""
        
        return html.encode('utf-8')
    
    async def _export_markdown(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to Markdown format"""
        lines = []
        
        lines.append(f"# {transcription.video_title or 'Transcripción de Video'}")
        lines.append("")
        
        if options.include_metadata:
            lines.append("## Información del Video")
            lines.append("")
            lines.append(f"- **Autor:** {transcription.video_author or 'Desconocido'}")
            lines.append(f"- **Plataforma:** {transcription.platform_detected.value if transcription.platform_detected else 'N/A'}")
            lines.append(f"- **Duración:** {self._format_duration(transcription.video_duration)}")
            lines.append(f"- **Fecha:** {transcription.created_at.strftime('%d/%m/%Y %H:%M')}")
            lines.append("")
        
        lines.append("## Transcripción")
        lines.append("")
        
        if options.include_segments and transcription.segments:
            for segment in transcription.segments:
                if options.include_timestamps:
                    lines.append(f"**{segment.formatted_timestamp}**")
                lines.append(f"> {segment.text}")
                lines.append("")
        else:
            lines.append(transcription.full_text or "")
        
        if options.include_analysis and transcription.analysis:
            lines.append("")
            lines.append("## Análisis de Contenido")
            lines.append("")
            lines.append(f"- **Framework:** {transcription.analysis.framework.value}")
            lines.append(f"- **Tono:** {transcription.analysis.tone}")
            lines.append(f"- **Audiencia:** {transcription.analysis.target_audience or 'General'}")
            
            if transcription.analysis.key_points:
                lines.append("")
                lines.append("### Puntos Clave")
                lines.append("")
                for point in transcription.analysis.key_points:
                    lines.append(f"- {point}")
            
            if transcription.analysis.hashtags_suggested:
                lines.append("")
                lines.append(f"**Hashtags:** {' '.join(transcription.analysis.hashtags_suggested)}")
        
        return '\n'.join(lines).encode('utf-8')
    
    async def _export_pdf(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to PDF format using reportlab"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=12,
                textColor=colors.HexColor('#1a1a2e')
            )
            
            story = []
            
            story.append(Paragraph(transcription.video_title or 'Transcripción', title_style))
            story.append(Spacer(1, 12))
            
            if options.include_metadata:
                meta_data = [
                    ['Autor:', transcription.video_author or 'Desconocido'],
                    ['Plataforma:', transcription.platform_detected.value if transcription.platform_detected else 'N/A'],
                    ['Duración:', self._format_duration(transcription.video_duration)],
                    ['Fecha:', transcription.created_at.strftime('%d/%m/%Y %H:%M')],
                ]
                meta_table = Table(meta_data, colWidths=[1.5*inch, 4*inch])
                meta_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#4a4e69')),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(meta_table)
                story.append(Spacer(1, 20))
            
            story.append(Paragraph('Transcripción', styles['Heading2']))
            story.append(Spacer(1, 12))
            
            text = transcription.full_text_with_timestamps if options.include_timestamps else transcription.full_text
            if text:
                for para in text.split('\n\n'):
                    if para.strip():
                        story.append(Paragraph(para.replace('\n', '<br/>'), styles['Normal']))
                        story.append(Spacer(1, 6))
            
            doc.build(story)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("reportlab not installed, falling back to HTML")
            html_content = await self._export_html(transcription, options)
            return html_content
    
    async def _export_docx(
        self,
        transcription: TranscriptionResponse,
        options: ExportOptions,
    ) -> bytes:
        """Export to DOCX format using python-docx"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            title = doc.add_heading(transcription.video_title or 'Transcripción', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if options.include_metadata:
                doc.add_paragraph()
                meta = doc.add_paragraph()
                meta.add_run('Autor: ').bold = True
                meta.add_run(transcription.video_author or 'Desconocido')
                
                meta = doc.add_paragraph()
                meta.add_run('Plataforma: ').bold = True
                meta.add_run(transcription.platform_detected.value if transcription.platform_detected else 'N/A')
                
                meta = doc.add_paragraph()
                meta.add_run('Duración: ').bold = True
                meta.add_run(self._format_duration(transcription.video_duration))
                
                meta = doc.add_paragraph()
                meta.add_run('Fecha: ').bold = True
                meta.add_run(transcription.created_at.strftime('%d/%m/%Y %H:%M'))
            
            doc.add_heading('Transcripción', level=1)
            
            if options.include_segments and transcription.segments:
                for segment in transcription.segments:
                    p = doc.add_paragraph()
                    if options.include_timestamps:
                        p.add_run(f'{segment.formatted_timestamp} ').bold = True
                    p.add_run(segment.text)
            else:
                doc.add_paragraph(transcription.full_text or '')
            
            if options.include_analysis and transcription.analysis:
                doc.add_heading('Análisis de Contenido', level=1)
                
                p = doc.add_paragraph()
                p.add_run('Framework: ').bold = True
                p.add_run(transcription.analysis.framework.value)
                
                p = doc.add_paragraph()
                p.add_run('Tono: ').bold = True
                p.add_run(transcription.analysis.tone)
                
                if transcription.analysis.key_points:
                    doc.add_heading('Puntos Clave', level=2)
                    for point in transcription.analysis.key_points:
                        doc.add_paragraph(point, style='List Bullet')
            
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
        except ImportError:
            logger.warning("python-docx not installed, falling back to HTML")
            html_content = await self._export_html(transcription, options)
            return html_content
    
    def _format_srt_time(self, seconds: float) -> str:
        """Format seconds to SRT time (HH:MM:SS,mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def _format_vtt_time(self, seconds: float) -> str:
        """Format seconds to VTT time (HH:MM:SS.mmm)"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    
    def _format_duration(self, seconds: Optional[float]) -> str:
        """Format duration for display"""
        if not seconds:
            return "N/A"
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        if minutes >= 60:
            hours = minutes // 60
            minutes = minutes % 60
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes}:{secs:02d}"


_export_service: Optional[ExportService] = None


def get_export_service() -> ExportService:
    """Get export service singleton"""
    global _export_service
    if _export_service is None:
        _export_service = ExportService()
    return _export_service












